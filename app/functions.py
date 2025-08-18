import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
import time
from app.prompts import INSTRUCOES_IA

load_dotenv()


def ler_conteudo_py(caminho_do_arquivo: str) -> str:
    """
    Lê o conteúdo de um arquivo Python (.py) e o retorna como uma única string,
    preservando toda a formatação original.

    Args:
        caminho_do_arquivo (str): O caminho completo para o arquivo .py que você deseja ler.

    Returns:
        str: O conteúdo do arquivo como uma string. Se o arquivo não for encontrado
             ou ocorrer outro erro, retorna uma mensagem de erro formatada.
    """
    if not caminho_do_arquivo.endswith('.py'):
        return f"Erro: O arquivo '{caminho_do_arquivo}' não parece ser um arquivo Python (.py)."

    try:
        with open(caminho_do_arquivo, 'r', encoding='utf-8') as arquivo:
            # mantendo todas as quebras de linha (\n) e espaços.
            conteudo_string = arquivo.read()
            return conteudo_string, os.path.basename(caminho_do_arquivo)
    except FileNotFoundError:
        return f"Erro: O arquivo no caminho '{caminho_do_arquivo}' não foi encontrado."
    except Exception as e:
        return f"Ocorreu um erro inesperado ao ler o arquivo: {e}"
    
    
def ler_conteudo_pas(caminho_do_arquivo: str):
    """
    Lê o conteúdo de um arquivo Pascal (.pas), tentando as codificações
    UTF-8 e latin-1, e o retorna como uma tupla contendo a string do
    conteúdo e o nome do arquivo.

    Args:
        caminho_do_arquivo (str): O caminho completo para o arquivo .pas.

    Returns:
        tuple[str, str] | str: Uma tupla (conteúdo, nome_arquivo) em caso de
                                 sucesso, ou uma string de erro em caso de falha.
    """
    if not caminho_do_arquivo.lower().endswith('.pas'):
        return f"Erro: O arquivo '{caminho_do_arquivo}' não parece ser um arquivo Pascal (.pas)."

    try:
        # 1ª TENTATIVA: Tenta ler com UTF-8, que é o padrão moderno
        with open(caminho_do_arquivo, 'r', encoding='utf-8') as arquivo:
            conteudo_string = arquivo.read()
            print("Arquivo lido com sucesso usando a codificação UTF-8.")
            return conteudo_string, os.path.basename(caminho_do_arquivo)
    except UnicodeDecodeError:
        # 2ª TENTATIVA: Se UTF-8 falhar, tenta com 'latin-1', comum em arquivos Windows
        print("UTF-8 falhou. Tentando ler com a codificação latin-1...")
        try:
            with open(caminho_do_arquivo, 'r', encoding='latin-1') as arquivo:
                conteudo_string = arquivo.read()
                print("Arquivo lido com sucesso usando a codificação latin-1.")
                return conteudo_string, os.path.basename(caminho_do_arquivo)
        except Exception as e:
            return f"Falha ao ler o arquivo com UTF-8 e também com latin-1. Erro final: {e}"
    except FileNotFoundError:
        return f"Erro: O arquivo no caminho '{caminho_do_arquivo}' não foi encontrado."
    except Exception as e:
        return f"Ocorreu um erro inesperado ao ler o arquivo: {e}"



def criar_docx_formatado(conteudo: str, nome_arquivo_final: str):
    """
    Cria um arquivo .docx com formatação específica
    """
    doc = Document()

    # Define a margem do documento
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Função auxiliar para adicionar texto com negrito
    def adicionar_texto_com_negrito(paragrafo, texto):
        partes = texto.split('**')
        for i, parte in enumerate(partes):
            if i % 2 == 1: # Partes ímpares estão em negrito
                paragrafo.add_run(parte).bold = True
            else:
                paragrafo.add_run(parte)

    # Processa o conteúdo linha por linha
    for linha in conteudo.strip().split('\n'):
        linha_strip = linha.strip()

        # Título Principal (H1)
        if linha_strip.startswith('H1: '):
            texto_titulo = linha_strip[4:]
            p = doc.add_paragraph()
            run = p.add_run(texto_titulo)
            run.font.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(12)
        
        # Títulos de Seção (H2)
        elif linha_strip.startswith('H2: '):
            texto_secao = linha_strip[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(texto_secao)
            run.font.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # Listas Aninhadas (com 2 espaços de indentação)
        elif linha_strip.startswith('  * '):
            texto_lista = linha_strip[4:]
            p = doc.add_paragraph(style='List Bullet 2')
            adicionar_texto_com_negrito(p, texto_lista)

        # Listas de Nível 1
        elif linha_strip.startswith('* '):
            texto_lista = linha_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            adicionar_texto_com_negrito(p, texto_lista)
        
        # Parágrafos Normais
        else:
            if linha_strip: # Adiciona apenas se a linha não estiver vazia
                p = doc.add_paragraph()
                adicionar_texto_com_negrito(p, linha_strip)
    
    if not nome_arquivo_final.lower().endswith('.docx'):
        nome_arquivo_final += '.docx'

    try:
        doc.save(f"docs/{nome_arquivo_final}")
        print(f"'{nome_arquivo_final}' salvo com sucesso")
    except Exception as e:
        print(f"Erro ao salvar o arquivo .docx: {e}")
        

def gerar_resposta_ia_document(codigo_para_analise: str, nome_do_arquivo: str, contexto_adicional: str = "Nenhum."):
    """
    Formata o prompt com o código e metadados, envia para a API do Google 
    e retorna a resposta da IA.

    Args:
        codigo_para_analise (str): O conteúdo do código Python a ser analisado.
        nome_do_arquivo (str): O nome do arquivo para usar na documentação.
        contexto_adicional (str, optional): Qualquer contexto extra a ser fornecido à IA. 
                                            Defaults to "Nenhum.".

    Returns:
        str: A documentação gerada pela IA (em str) ou uma mensagem de erro.
    """
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            temperature=0.1,                # Temperatura muito baixa para alta fidelidade ao código
            google_api_key=os.getenv("GEMINI_API_KEY")
        )

        # Usa o PromptTemplate para preencher as variáveis do seu arquivo de prompts
        prompt_template = PromptTemplate.from_template(INSTRUCOES_IA)

        # Cria a "cadeia" de execução: o prompt formatado é enviado para o LLM
        chain = prompt_template | llm

        print(f"Analisando o arquivo '{nome_do_arquivo}'")
        
        start_time = time.time()
        # Invoca a cadeia, passando os valores para os placeholders do seu prompt
        resposta_ia = chain.invoke({
            "nome_arquivo": nome_do_arquivo,
            "contexto_extra": contexto_adicional,
            "codigo": codigo_para_analise
        })

        end_time = time.time()
        tempo_total = round(end_time - start_time, 2)
        print(f"Resposta da IA recebida em {tempo_total} segundos.")
        
        # A resposta do modelo fica no atributo .content
        return resposta_ia.content

    except Exception as e:
        print(f"Erro ao comunicar com a IA: {e}")
        return f"Erro ao gerar resposta da IA: {e}"
