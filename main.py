import os
import sys
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# --- INICIALIZAÇÃO E CONFIGURAÇÃO ---

# Usando sua função importada
# Certifique-se de que o arquivo 'functions.py' está no mesmo diretório ou no PYTHONPATH
try:
    from functions import ler_conteudo_py
except ImportError:
    print("❌ Erro: Não foi possível encontrar o arquivo 'functions.py'.")
    print("   Certifique-se de que ele está no mesmo diretório que este script.")
    sys.exit(1)

load_dotenv()

# Caminho do arquivo a ser documentado está definido diretamente no código.
caminho_do_arquivo_alvo = "C:\\Users\\marcos.rodrigues\\Documents\\documentation\\doc_bic_distribui_xml.py"
print(f"📄 Lendo o arquivo alvo: {caminho_do_arquivo_alvo}")


# Chama sua função importada para ler o arquivo
str_codigo_fonte, nome_arquivo = ler_conteudo_py(caminho_do_arquivo_alvo)

# Valida se a leitura do arquivo foi bem-sucedida
if not str_codigo_fonte or not nome_arquivo:
    print(f"❌ Erro ao ler o arquivo: {caminho_do_arquivo_alvo}. Verifique o caminho e a função 'ler_conteudo_py'.")
    sys.exit(1)

# Sanitiza o código-fonte para evitar conflitos com a formatação do prompt da IA
str_codigo_safe = str_codigo_fonte.replace("```", "'''").replace('"""', "'''")


# --- PROMPT PARA A INTELIGÊNCIA ARTIFICIAL (COM EXEMPLOS COMPLETOS E REGRAS RÍGIDAS) ---

INSTRUCOES_IA = f"""
Você é um assistente de IA especialista em criar documentação técnica detalhada e factual para scripts Python.

**INSTRUÇÃO PRINCIPAL:**
Sua tarefa é gerar uma documentação para o script fornecido. Estude os **três exemplos completos e de alta qualidade** abaixo. Eles são seu **guia de estilo, tom e, mais importante, NÍVEL DE DETALHE**. A estrutura (as seções H2) deve ser **adaptada** ao script que você está analisando.

**REGRAS ANTI-ALUCINAÇÃO (MUITO IMPORTANTE):**
- **BASEIE-SE APENAS NO CÓDIGO-FONTE FORNECIDO.** Não invente funcionalidades, bibliotecas, nomes de arquivos, parâmetros ou lógicas que não estão explicitamente no código.
- **SE UMA INFORMAÇÃO NÃO ESTÁ NO CÓDIGO, NÃO A ADICIONE.** É melhor omitir um detalhe do que inventar um. Se não conseguir determinar uma informação (como 'Frequência' ou 'Servidor'), declare isso explicitamente. Exemplo: "A frequência de execução não pode ser determinada a partir do código-fonte."
- **FOQUE EM DESCREVER O QUE EXISTE.** Sua tarefa é documentar o script como ele é, não como ele poderia ser. Não faça suposições sobre o ambiente ou sobre módulos externos (como 'tools').

**REGRAS DE DETALHAMENTO:**
- **VÁ ALÉM DO ÓBVIO:** Explique o **propósito de negócio** de cada etapa, as **decisões de design** (por que foi feito assim?) e como as partes se conectam.
- **SEJA TÉCNICO E PRECISO:** Use os nomes corretos de variáveis, funções e tabelas mencionados no código.
- **ADAPTE A ESTRUTURA:** Omita seções dos exemplos que não se aplicam e adicione novas seções se o script tiver funcionalidades únicas que mereçam destaque.

**REGRAS DE FORMATAÇÃO OBRIGATÓRIAS:**
- Título principal: `H1: `
- Títulos de seção: `H2: `
- Listas: `* ` (use `  * ` para listas aninhadas).
- Negrito: `**este exemplo**`.

---
**EXEMPLO DE DOCUMENTAÇÃO 1 (FOCO EM HEALTH CHECK DE BANCO DE DADOS):**

H1: Documentação – alerta_tracking.py

H2: Área relacionada
* Logix

H2: Frequência
* O script é executado a cada 15 minutos, todos os dias.

H2: Servidor
* 192.168.102.169

H2: Resumo da aplicação
Esta aplicação consiste em um script Python (alerta_tracking.py) que funciona como um monitor de saúde (health check) para o banco de dados do sistema de Tracking. Seu único objetivo é verificar se o banco de dados TRACKING está online e respondendo a conexões. Para isso, ele tenta se conectar e executar uma consulta extremamente simples. Se a conexão ou a consulta falharem, o script assume que o banco de dados está com problemas e imediatamente dispara um alerta. O alerta é inserido na tabela alerta_mensagens (no banco TORRE_CONTROLE) para notificar a equipe de TI responsável sobre a indisponibilidade.

H2: Pré-requisitos
* **Bibliotecas Python:**
  * **Não Nativas:** python-dotenv. Nota: Requer um driver de banco de dados (como pyodbc ou psycopg2), que é utilizado dentro do módulo 'tools'.
  * **Nativas:** os, sys, datetime.
* **Arquivos Necessários:**
  * Módulo 'tools' acessível via path.
  * Um arquivo .env no diretório de ferramentas para armazenar as credenciais dos bancos de dados.
* **Parâmetros de Linha de Comando:** O script precisa ser executado com parâmetros que definem o canal do alerta e os destinatários: `[template] [destinatario1] [destinatario2] ...`

H2: Fluxo de Execução Detalhado
**Inicialização e Leitura de Parâmetros:** O script define os caminhos de sistema, carrega as variáveis de ambiente e lê os parâmetros passados via linha de comando para determinar o modelo da mensagem e a lista de contatos que devem ser notificados em caso de falha.
**Verificação de Saúde do Banco de Dados (Health Check):** Esta é a etapa central. O script tenta se conectar ao banco de dados TRACKING. Se a conexão for bem-sucedida, ele executa uma consulta mínima (`SELECT 1;`), que não busca dados reais, mas serve apenas para confirmar que o banco está funcional e apto a processar requisições. Toda essa operação é envolvida por um bloco `try...except` para capturar qualquer tipo de erro de conexão ou execução.
**Análise do Resultado:** Uma falha é determinada por duas condições: 1) Se ocorrer qualquer exceção durante a tentativa de conexão ou de execução da consulta. 2) Se a consulta for executada com sucesso, mas não retornar o resultado esperado (uma linha), indicando um comportamento anômalo. Se qualquer uma dessas condições for verdadeira, uma flag interna de erro é ativada.
**Tratamento de Sucesso:** Se, após a verificação, a flag de erro permanecer desativada, o script considera que o banco de dados TRACKING está saudável e encerra sua execução imediatamente, sem realizar nenhuma ação.
**Geração de Alerta em Caso de Falha:** Se a flag de erro for ativada, o script entra no modo de alerta. Ele constrói uma mensagem padrão informando qual banco de dados (TRACKING) parou de responder e o horário da ocorrência.
**Enfileiramento do Alerta:** O script então se conecta a um segundo banco de dados, o TORRE_CONTROLE. Ele percorre a lista de destinatários recebida como parâmetro e, para cada um, insere uma nova linha na tabela `alerta_mensagens`. Essa linha contém a mensagem de falha e os detalhes para o envio, efetivamente "enfileirando" o alerta para ser despachado por um sistema de notificação separado.
**Finalização:** Após registrar todos os alertas, a conexão com o banco TORRE_CONTROLE é fechada, e o script termina.

H2: Funções Auxiliares Principais
A lógica do script é executada de forma linear em um fluxo único e não define funções auxiliares próprias. Ele depende inteiramente de funções do módulo externo 'tools' para tarefas padronizadas, como a conexão com os bancos de dados (t.connect_db), o fechamento dessas conexões (t.close_db) e o registro de logs (t.log).

---
**EXEMPLO DE DOCUMENTAÇÃO 2 (FOCO EM MONITORAMENTO DE API):**

H1: Documentação – alerta_usr_api_wms.py

H2: Área relacionada
* Logix

H2: Frequência
* O script é executado a cada 10 minutos, todos os dias.

H2: Servidor
* 192.168.102.169

H2: Resumo da aplicação
Esta aplicação consiste em um script Python (alerta_usr_api_wms.py) que funciona como uma ferramenta de monitoramento proativo. Seu objetivo é verificar continuamente se as credenciais de um usuário de sistema, utilizado para integrações via API com o Protheus/WMS, estão ativas e funcionando corretamente. O script opera no modelo "silêncio em caso de sucesso": ele tenta se autenticar na API e, se conseguir, não faz nada. Caso a autenticação falhe por qualquer motivo (senha expirada, usuário bloqueado, API fora do ar), ele imediatamente gera um alerta. Esse alerta é inserido em uma tabela de controle (alerta_mensagens) para ser enviado a uma lista de destinatários, garantindo que a equipe responsável seja notificada rapidamente sobre a falha.

H2: Pré-requisitos
* **Bibliotecas Python:**
  * **Não Nativas:** requests, python-dotenv.
  * **Nativas:** os, sys, datetime, json.
* **Arquivos Necessários:**
  * Módulo 'tools' acessível via pathtools.
  * Um arquivo .env para armazenar as credenciais da API no formato `usr_api_[identificador]` e `pwd_api_[identificador]`.
* **Parâmetros de Linha de Comando:** `[canal_de_alerta] [identificador_do_usuario] [destinatario1] [destinatario2] ...`

H2: Fluxo de Execução Detalhado
**Inicialização e Leitura de Parâmetros:** O script define os caminhos de sistema, carrega as variáveis de ambiente e lê os parâmetros da linha de comando para determinar o identificador do usuário a ser testado (ex: 'canadian'), o canal de alerta (ex: 'WA') e a lista de destinatários.
**Busca de Credenciais:** Utilizando o identificador, o script busca as credenciais correspondentes (usuário e senha) no arquivo .env.
**Tentativa de Autenticação na API:** O script envia uma requisição POST para o endpoint de autenticação da API do WMS ([https://api.supptech.com.br/wms/api/oauth2/v1/token](https://api.supptech.com.br/wms/api/oauth2/v1/token)).
**Tratamento de Sucesso:** Se a API retornar o código de status 201 (Created), a autenticação foi bem-sucedida, e o script encerra silenciosamente.
**Geração de Alerta em Caso de Falha:** Se o código de status for diferente de 201, o script trata como uma falha, construindo uma mensagem de erro detalhada.
**Enfileiramento do Alerta:** O script conecta-se ao banco TORRE_CONTROLE e insere uma nova linha na tabela `alerta_mensagens` para cada destinatário, enfileirando os alertas para envio.
**Finalização:** A conexão com o banco é fechada e o script termina.

H2: Funções Auxiliares Principais
* **main():** Concentra todo o fluxo de trabalho.
* O script utiliza funções do módulo externo 'tools' para tarefas padronizadas (conexão com DB, logs, etc.).

---
**EXEMPLO DE DOCUMENTAÇÃO 3 (FOCO EM PROCESSAMENTO DE E-MAILS):**

H1: Documentação – nfs_captura_anexos_remove.py

H2: Área relacionada
* Logix

H2: Frequência
* O script é executado a cada 5 minutos, todos os dias.

H2: Servidor
* 192.168.102.169

H2: Resumo da aplicação
Esta aplicação consiste em um script Python (nfs_captura_anexos_remove.py) que funciona como um processador de e-mails automatizado para a caixa de entrada nfs.parceiros@supplog.com. O script lê os e-mails não lidos, identifica a qual cliente parceiro o e-mail pertence com base em palavras-chave no assunto e, em seguida, extrai os anexos (XML, PDF e ZIP). Os arquivos XML são registrados em uma tabela de controle no banco de dados (auto_nfs) e salvos em pastas de rede específicas para serem importados pelo sistema Logix, enquanto os PDFs são salvos para backup. Após o processamento, os e-mails são movidos ou deletados para manter a caixa de entrada organizada.

H2: Pré-requisitos
* **Bibliotecas Python:**
  * **Não Nativas:** python-dotenv. Nota: O script depende de bibliotecas para comunicação com APIs (como requests), que são utilizadas dentro do módulo tools.
  * **Nativas:** os, sys, datetime, pathlib, zipfile, shutil, email.
* **Arquivos Necessários:**
  * Módulo 'tools' acessível via pathtools de acordo com o SO.
  * Um arquivo .env no diretório de ferramentas para armazenar as variáveis de ambiente (credenciais do banco e da API de e-mail).

H2: Fluxo de Execução Detalhado
**Inicialização e Configuração:** O script define os caminhos de sistema, carrega as variáveis de ambiente e cria um conjunto de dicionários que funcionam como um mapa de configuração. Esses dicionários associam cada cliente a palavras-chave específicas de assunto de e-mail e aos seus respectivos caminhos de pasta para salvar os arquivos XML e PDF.
**Autenticação e Leitura de E-mails:** O script obtém um token de autenticação para a conta nfs.parceiros@supplog.com e o utiliza para buscar até 500 e-mails que estejam marcados como "não lidos" na caixa de entrada.
**Processamento e Classificação de E-mails:** O script inicia um loop, processando cada e-mail individualmente. A primeira etapa é usar o assunto do e-mail para identificar o cliente correspondente, consultando o dicionário de configuração. Se nenhum cliente for identificado, o e-mail é considerado irrelevante.
**Extração e Tratamento de Anexos:** Se um cliente é identificado, o script baixa todos os anexos do e-mail. Ele trata cada tipo de anexo de forma diferente:
  * **XML:** O script primeiro tenta inserir os dados da nota (cliente, número, série) na tabela de controle auto_nfs. Se o registro for novo (não um duplicado), o arquivo XML é salvo em duas pastas: uma para o Logix importar e outra para backup.
  * **PDF:** O arquivo é salvo diretamente na pasta de backup do cliente.
  * **ZIP:** O arquivo .zip é salvo em uma pasta temporária e seu conteúdo é extraído. O script então processa os arquivos de dentro do ZIP (XMLs e PDFs) da mesma forma que faria com anexos normais.
  * **E-mail Anexado (.eml):** Se um anexo for outro e-mail, o script o abre e processa os anexos contidos nele de forma recursiva.
**Limpeza da Caixa de Entrada:** Após processar um e-mail, uma ação de limpeza é executada:
  * **E-mails de Clientes Conhecidos:** São marcados como "lidos" e movidos para a pasta "Itens Excluídos".
  * **E-mails de Clientes Não Identificados:** São marcados como "lidos" e deletados permanentemente, pois não são relevantes para o processo.
**Finalização:** Após o loop por todos os e-mails, a conexão com o banco de dados é fechada para liberar os recursos.

H2: Funções Auxiliares Principais
* **processa():** É a função principal que orquestra todo o fluxo de trabalho.
* **insere_auto_nfs(cliente, nome_arquivo, ...):** Função crucial responsável por interagir com o banco de dados. Utiliza a cláusula ON CONFLICT DO NOTHING do PostgreSQL para evitar a inserção de notas duplicadas.
* **salva_anexos_XML_PDF_LOGIX(cliente, eml, ...):** Uma função especializada para tratar anexos .eml.
* **find_key_with_string(dicionario, texto):** Utilitário de busca que encontra o cliente com base no assunto do e-mail.
* O script depende fortemente de funções do módulo externo **tools**.

H2: Lógica por Cliente
O script utiliza dicionários para gerenciar as regras de cada cliente.
* **d_assunto:** Mapeia o nome do cliente a strings de identificação no assunto.
* **d_pasta_logix:** Define o diretório de entrada para o Logix.
* **d_pasta_xml e d_pasta_pdf:** Definem os diretórios de backup.
* **armor:** Identificação: "FATURAMENTO DE UBERLANDIA DO DIA ".
* **vitoria_regia:** Identificação: "INDUSTRIAS REUNIDAS VITORIA REGIA LTDA.: NF-e Nº " ou "FATURAMENTO VITORIA RÉGIA - ".
* **dbs:** Identificação: "Envio". Regra Especial: extrai todos os caracteres numéricos do nome do arquivo para formar o número da nota.
* **ddw:** Identificação: "ENC: RESUMO DE FATURAMENTO - " ou "RESUMO DE FATURAMENTO - ".
* **latin_master:** Identificação: "RESUMO DE FATURAMENTO LATIN MASTER".
* **epa:** Identificação: "OPERAÇÃO IPOJUCA CLIENTE:".

H2: Cliente não identificado
Se o assunto de um e-mail não corresponder a nenhuma palavra-chave, o script o marca como "lido" e o deleta permanentemente.
---

**SUA TAREFA AGORA:**

Baseado em **TODAS as regras acima** e nos **três exemplos completos e de alta qualidade**, analise as informações a seguir e gere a documentação mais precisa, detalhada e factual possível para o script fornecido.

**INFORMAÇÕES PARA ANÁLISE:**

**1. NOME DO SCRIPT:**
{nome_arquivo}

**2. CÓDIGO-FONTE DO SCRIPT PYTHON:**
```python
{str_codigo_safe}
```

**TAREFA FINAL:**
Gere a documentação completa, detalhada e adaptada para o script `{nome_arquivo}`.
"""

# Modelo de IA e nome do arquivo de saída
modelo_ia = "llama3-70b-8192"
nome_arquivo_saida = f"DOC_{nome_arquivo.split('.')[0]}.docx"


# --- FUNÇÕES PRINCIPAIS ---

def gerar_conteudo_com_ia(instrucoes: str, modelo: str) -> str:
    """
    Envia o prompt para a API da Groq e retorna a resposta da IA.
    """
    print("🤖 Analisando o código e gerando a documentação com a IA (usando múltiplos exemplos e regras anti-alucinação)...")
    try:
        api_key = os.environ.get("GROQ_API_KEY")
        if not api_key:
            raise ValueError("A variável de ambiente GROQ_API_KEY não foi definida. Verifique seu arquivo .env.")
            
        client = Groq(api_key=api_key)
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": instrucoes,
                },
                {
                    "role": "user",
                    "content": f"Por favor, gere a documentação completa e adaptada para o script {nome_arquivo} conforme solicitado nas instruções.",
                },
            ],
            model=modelo,
            temperature=0.0, # Temperatura zerada para máxima aderência aos fatos e mínima criatividade
        )
        print("✅ Documentação gerada com sucesso!")
        return chat_completion.choices[0].message.content
    except Exception as e:
        print(f"❌ Erro ao se comunicar com a API da Groq: {e}")
        return ""

def criar_docx_formatado(conteudo: str, nome_arquivo_final: str):
    """
    Cria um arquivo .docx com formatação específica, replicando o estilo do documento de exemplo.
    """
    print(f"📄 Criando o arquivo .docx '{nome_arquivo_final}' com formatação avançada...")
    doc = Document()
    doc.core_properties.title = "Documentação de Script"
    doc.core_properties.author = "Gerador de Documentação com IA"

    # Define a margem do documento
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Função auxiliar para adicionar texto com negrito
    def adicionar_texto_com_negrito(paragrafo, texto):
        partes = texto.split('**')
        for i, parte in enumerate(partes):
            if i % 2 == 1: # Partes ímpares estão em negrito
                paragrafo.add_run(parte).bold = True
            else:
                paragrafo.add_run(parte)

    # Processa o conteúdo linha por linha
    for linha in conteudo.strip().split('\n'):
        linha_strip = linha.strip()

        # Título Principal (H1)
        if linha_strip.startswith('H1: '):
            texto_titulo = linha_strip[4:]
            p = doc.add_paragraph()
            run = p.add_run(texto_titulo)
            run.font.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(12)
        
        # Títulos de Seção (H2)
        elif linha_strip.startswith('H2: '):
            texto_secao = linha_strip[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(texto_secao)
            run.font.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # Listas Aninhadas (com 2 espaços de indentação)
        elif linha_strip.startswith('  * '):
            texto_lista = linha_strip[4:]
            p = doc.add_paragraph(style='List Bullet 2')
            adicionar_texto_com_negrito(p, texto_lista)

        # Listas de Nível 1
        elif linha_strip.startswith('* '):
            texto_lista = linha_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            adicionar_texto_com_negrito(p, texto_lista)
        
        # Parágrafos Normais
        else:
            if linha_strip: # Adiciona apenas se a linha não estiver vazia
                p = doc.add_paragraph()
                adicionar_texto_com_negrito(p, linha_strip)

    try:
        doc.save(nome_arquivo_final)
        print(f"✅ Documento '{nome_arquivo_final}' salvo com sucesso!")
    except Exception as e:
        print(f"❌ Erro ao salvar o arquivo .docx: {e}")


# --- BLOCO DE EXECUÇÃO PRINCIPAL ---

if __name__ == "__main__":
    conteudo_gerado = gerar_conteudo_com_ia(INSTRUCOES_IA, modelo_ia)
    
    if conteudo_gerado:
        criar_docx_formatado(conteudo_gerado, nome_arquivo_saida)
